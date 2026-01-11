"""
DOCX formatter for quizzes with rubrics.

Generates Word documents with:
- Quiz header (subject, topic, duration, CLOs)
- Questions with CLO alignment and Bloom's level
- Detailed rubrics for each question
- Professional formatting
"""

from pathlib import Path
from typing import Any, Optional
from datetime import datetime

from shared.formatters.base_formatter import BaseFormatter
from shared.utils.metadata_manager import create_quiz_metadata

# Conditional import for python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxFormatter(BaseFormatter):
    """Formatter for generating Word document quizzes."""

    @property
    def material_type(self) -> str:
        return "quizzes"

    @property
    def output_format(self) -> str:
        return "docx"

    def generate(self, content: dict[str, Any]) -> Path:
        """
        Generate Word document quiz.

        Args:
            content: Dictionary containing:
                - version: str
                - time_duration: int (minutes)
                - total_questions: int
                - complexity_levels: list[str]
                - question_types: list[str]
                - clos: list[str] (Course Learning Outcomes)
                - questions: list[dict] with:
                    - number: int
                    - text: str
                    - clo_number: int
                    - bloom_level: str
                    - marks: int
                    - rubric: dict with criteria and performance levels
                - reference: dict (optional)

        Returns:
            Path to the generated Word document
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            )

        self.ensure_output_directory()
        output_path = self.get_output_path()

        # Create document
        doc = Document()
        self._setup_styles(doc)

        # Header
        self._add_header(doc, content)

        # CLOs section
        self._add_clos_section(doc, content)

        # Instructions
        self._add_instructions(doc, content)

        # Questions with rubrics
        for question in content.get("questions", []):
            self._add_question(doc, question)
            self._add_rubric(doc, question)

        # Save document
        doc.save(str(output_path))

        # Save metadata
        metadata = create_quiz_metadata(
            topic=self.topic,
            subject=self.subject,
            clos=content.get("clos", []),
            time_duration=content.get("time_duration", 60),
            total_questions=content.get("total_questions", 5),
            complexity_levels=content.get("complexity_levels", ["Apply"]),
            question_types=content.get("question_types", ["Short Answer"]),
            reference=content.get("reference")
        )
        metadata["current_version"] = content.get("version", "v1.0")
        self.save_metadata(metadata)

        return output_path

    def _setup_styles(self, doc: Document):
        """Setup custom document styles."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_header(self, doc: Document, content: dict[str, Any]):
        """Add quiz header section."""
        version = content.get("version", "v1.0")
        duration = content.get("time_duration", 60)
        total_questions = content.get("total_questions", 5)

        # Title
        title = doc.add_heading(self.topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Assessment Quiz")
        run.font.size = Pt(14)
        run.font.italic = True

        doc.add_paragraph()

        # Metadata table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        metadata_rows = [
            ("Subject:", self.subject),
            ("Topic:", self.topic),
            ("Duration:", f"{duration} minutes"),
            ("Questions:", str(total_questions)),
            ("Version:", self._format_version_header(version)),
        ]

        for i, (label, value) in enumerate(metadata_rows):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

    def _add_clos_section(self, doc: Document, content: dict[str, Any]):
        """Add Course Learning Outcomes section."""
        clos = content.get("clos", [])
        if not clos:
            return

        doc.add_heading("Course Learning Outcomes (CLOs)", level=1)

        for i, clo in enumerate(clos, 1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(f"CLO {i}: ").bold = True
            para.add_run(clo)

        doc.add_paragraph()

    def _add_instructions(self, doc: Document, content: dict[str, Any]):
        """Add quiz instructions section."""
        duration = content.get("time_duration", 60)

        doc.add_heading("Instructions", level=1)

        instructions = [
            "Read all questions carefully before answering.",
            "Answer all questions.",
            f"Time limit: {duration} minutes.",
            "Refer to the rubric for grading criteria.",
            "Show all work for problem-solving questions.",
        ]

        for instruction in instructions:
            doc.add_paragraph(instruction, style='List Bullet')

        doc.add_paragraph()
        doc.add_page_break()

    def _add_question(self, doc: Document, question: dict[str, Any]):
        """Add a question to the document."""
        number = question.get("number", 1)
        text = question.get("text", "")
        clo_number = question.get("clo_number", 1)
        bloom_level = question.get("bloom_level", "Apply")
        marks = question.get("marks", 10)

        # Question heading
        heading = doc.add_heading(f"Question {number}", level=2)

        # Question metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"CLO Alignment: ").bold = True
        meta_para.add_run(f"CLO #{clo_number}  |  ")
        meta_para.add_run(f"Bloom's Level: ").bold = True
        meta_para.add_run(f"{bloom_level}  |  ")
        meta_para.add_run(f"Marks: ").bold = True
        meta_para.add_run(str(marks))

        # Question text
        doc.add_paragraph()
        question_para = doc.add_paragraph(text)
        question_para.paragraph_format.left_indent = Inches(0.25)

        doc.add_paragraph()

    def _add_rubric(self, doc: Document, question: dict[str, Any]):
        """Add rubric for a question."""
        number = question.get("number", 1)
        rubric = question.get("rubric", {})

        if not rubric:
            return

        # Rubric heading
        doc.add_heading(f"Rubric for Question {number}", level=3)

        # Criteria table
        criteria = rubric.get("criteria", [])
        if criteria:
            doc.add_paragraph().add_run("Criteria:").bold = True

            table = doc.add_table(rows=len(criteria) + 1, cols=3)
            table.style = 'Table Grid'

            # Header row
            header_row = table.rows[0]
            header_row.cells[0].text = "#"
            header_row.cells[1].text = "Criterion"
            header_row.cells[2].text = "Marks"
            for cell in header_row.cells:
                cell.paragraphs[0].runs[0].font.bold = True

            # Criteria rows
            for i, criterion in enumerate(criteria, 1):
                row = table.rows[i]
                row.cells[0].text = str(i)
                row.cells[1].text = criterion.get("description", "")
                row.cells[2].text = str(criterion.get("marks", 0))

        doc.add_paragraph()

        # Performance levels
        levels = rubric.get("performance_levels", {})
        if levels:
            doc.add_paragraph().add_run("Performance Levels:").bold = True

            level_order = ["Excellent", "Good", "Satisfactory", "Needs Improvement"]
            level_ranges = {
                "Excellent": "90-100%",
                "Good": "75-89%",
                "Satisfactory": "60-74%",
                "Needs Improvement": "<60%"
            }

            for level in level_order:
                if level in levels:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"{level} ({level_ranges.get(level, '')}): ").bold = True
                    para.add_run(levels[level])

        doc.add_paragraph()
        doc.add_paragraph("---")  # Separator
        doc.add_paragraph()
